from flask import Flask, render_template, request
from docx import Document
import re

app = Flask(__name__)

def convert_docx_to_html(file_stream):
    """
    Convert DOCX file to HTML-like text with <br> for line breaks.
    Strips any inline styles to avoid forcing fonts in the input box.
    """
    doc = Document(file_stream)
    full_text = ""

    # Add paragraphs
    for para in doc.paragraphs:
        full_text += para.text + "\n\n"

    # Add tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    full_text += para.text + "\n\n"

    # Convert newlines to <br> for HTML display
    html_text = full_text.replace("\n", "<br>")

    # Strip any style attributes (safety)
    html_text = re.sub(r'style="[^"]*"', '', html_text)

    return html_text

@app.route("/", methods=["GET", "POST"])
def index():
    converted_text = ""  # default empty input
    if request.method == "POST":
        uploaded_file = request.files.get("docx_file")
        if uploaded_file:
            converted_text = convert_docx_to_html(uploaded_file.stream)

    return render_template("index.html", converted_text=converted_text)

if __name__ == "__main__":
    app.run(debug=True)